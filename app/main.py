import os
import re
import torch
import uuid
import datetime
import torch.nn.functional as F
from markdown import markdown
from weasyprint import HTML
from openai import OpenAI
from fastapi import FastAPI, Body
from pydantic import BaseModel, Field
from docx import Document
from fastapi.responses import FileResponse
from transformers import AutoTokenizer, AutoModelForSequenceClassification

# FastAPI 앱 생성(Swagger 문서 정보 포함함)
app = FastAPI(
    title = 'AI 기반 코드 보안 분석기 API',
    description = "CodeBERT와 GPT API를 활용해 코드의 보안 취약점을 자동 탐지하고, 보안 리포트를 생성하는 API입니다.",
    version = "1.0.0",
    contact = {
        "email": "davin0706@gmail.com",
        "url": "https://github.com/tealight03/CapstoneDesign"
    }
)

# 입력 스키마
class CodeRequest(BaseModel):
    code: str = Field(
        ...,
        description = "보안 분석할 소스 코드 (Python)",
        example = "user_input = input('Enter ID: ')\nsql = 'SELECT * FROM users WHERE id = ' + user_input"
    )
    
# 응답 스키마
class AnalyzeResponse(BaseModel):
    prediction: str = Field(..., description="분석 결과 메시지")
    label: str = Field(..., description="예측된 보안 취약점 라벨 (예: SQL_Injection, Safe_Code 등)")
    security_score: int = Field(..., description="0~100 사이의 보안 점수 (낮을수록 취약함)")
    report: str = Field(..., description="GPT API가 작성한 보안 분석 보고서")
    model_reference: dict = Field(..., description="CodeBERT 모델의 원래 예측 결과")

# 모델 경로를 Hugging Face 경로로 설정
model_path = "davin0706/codebert-finetuned-v5"
tokenizer = AutoTokenizer.from_pretrained(model_path)
model = AutoModelForSequenceClassification.from_pretrained(model_path)

# ✅ OpenAI 클라이언트 초기화
client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

# ✅ 점수 매핑
score_map = {
    0: {"label": "SQL_Injection", "score": 30, "msg": "⚠️ SQL Injection 취약점 감지"},
    1: {"label": "Hardcoded_Password", "score": 30, "msg": "⚠️ Hardcoded Password 감지"},
    2: {"label": "XSS", "score": 30, "msg": "⚠️ XSS 취약점 감지"},
    3: {"label": "Safe_Code", "score": 100, "msg": "✅ 완전한 방어 코드 (Safe Code)"},
    4: {"label": "Other", "score": 70, "msg": "⚠️ 잠재적 취약점이 있는 코드 (Other)"}
}

# 라벨 확인 함수
def resolve_final_label_score(bert_label: str, gpt_label: str = None) -> dict:
    print(f"🔍 [라벨 결정] CodeBERT 라벨: '{bert_label}', GPT 라벨: '{gpt_label}'")
    final_entry = None

    # 1. GPT 라벨이 존재하고 CodeBERT와 다르면 우선 적용 시도
    if gpt_label and gpt_label.lower() != bert_label.lower():
        for entry in score_map.values():
            if entry["label"].lower() == gpt_label.lower():
                final_entry = entry
                print(f"✅ [GPT 라벨 채택] '{gpt_label}'에 해당하는 라벨 정보 적용")
                break
        if not final_entry:
            print(f"⚠️ [GPT 라벨 미일치] GPT 라벨 '{gpt_label}'이 유효하지 않음. CodeBERT 결과로 fallback")

    # 2. fallback: CodeBERT 라벨 기반
    if not final_entry:
        for entry in score_map.values():
            if entry["label"].lower() == bert_label.lower():
                final_entry = entry
                print(f"🔁 [Fallback] CodeBERT 라벨 '{bert_label}' 기반 결과 적용")
                break

    if final_entry:
        print(f"🏁 [최종 선택] 라벨: {final_entry['label']}, 점수: {final_entry['score']}, 메시지: {final_entry['msg']}")
    else:
        print("❌ [에러] 라벨 매핑 실패: 기본값 반환 예정")

    return final_entry


# ✅ 보안 분석 함수
def analyze_code(code_snippet: str):
    inputs = tokenizer(code_snippet, return_tensors="pt", truncation=True, padding="max_length", max_length=128)
    with torch.no_grad():
        outputs = model(**inputs)
        logits = outputs.logits
        probs = F.softmax(logits, dim=1)[0]
        predicted = torch.argmax(probs).item()

        # Safe vs Other 후처리
        if predicted == 3:
            if probs[4].item() > 0.3 and probs[3].item() < 0.7:
                predicted = 4

    return {
        "prediction": score_map[predicted]["msg"],
        "label": score_map[predicted]["label"],
        "security_score": score_map[predicted]["score"]
    }

# ✅ GPT 리포트 생성 함수
def generate_report(code_snippet: str, label: str) -> str:
    prompt = f"""
당신은 보안 분석 전문가입니다.
모델 분석 결과, 아래의 소스 코드에서 감지된 취약점은 **'{label}'** 입니다.
아래의 소스 코드를 상세히 분석해서 예상되는 취약점 보고서를 작성해주세요.

[취약 코드]
{code_snippet.rstrip()}


📌 **1. 취약점 설명**
해당 취약점이 발생한 이유와 코드 구조상 문제점을 설명해주세요.
"해당 코드는 ~~~ 취약점이 있는 것으로 분석되었습니다." 와 같은 문장으로 문장을 시작해주세요.

💣 **2. 공격 시나리오**
공격자가 해당 코드를 어떻게 악용할 수 있는지 설명해주세요.

🛠 **3. 보완 방법 및 개선된 코드 예시**
보완 방법을 설명하고, 보완된 코드 예시를 포함해주세요.

✅ **4. 요약된 보안 권고사항**
요약된 권고사항을 **리스트 형식**으로 정리해주세요.

---

아래 형식으로 리포트를 작성해주세요:

📄 **보안 분석 리포트**

---
"""
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant for code security analysis."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
        max_tokens=1200
    )
    return response.choices[0].message.content

# GPT API 결과에서 label 추출
def extract_label_from_report(report: str) -> str:
    try:
        # 정규식으로 취약점 라벨 추출 시도
        match = re.search(r"취약점(?:은|이)?[^가-힣]{0,5}\*\*['\"]?(.+?)['\"]?\*\*", report)
        if match:
            candidate = match.group(1).strip().replace(" ", "_")
            print(f"🔍 [정규식 매칭] 추출된 candidate: '{candidate}'")

            # 안전한 코드로 판명되었을 경우
            if candidate == "Safe_Code":
                for item in score_map.values():
                    if item["label"] != "Safe_Code" and item["label"] in report:
                        print(f"⚠️ Safe_Code 오탐 감지, 대체 라벨 사용: '{item['label']}'")
                        return item["label"]
            print(f"✅ 최종 선택된 라벨 (정규식 기준): '{candidate}'")
            return candidate

        # fallback: 정규식 실패 시, 본문 내에서 수동 탐색
        for item in score_map.values():
            if item["label"] in report:
                print(f"🔁 [본문 탐색] '{item['label']}' 라벨이 보고서에 직접 포함되어 있음")
                return item["label"]

    except Exception as e:
        print(f"⚠️ extract_label_from_report 예외 발생: {e}")
    
    print("❌ 라벨 추출 실패: None 반환")
    return None

# / 를 서버 헬스 체크 엔드포인트로 활용
@app.get("/", summary="API 서버 헬스 체크")
def root():
    return {"message": "🚀 Code Security Analyzer is running!"}

# ✅ API 엔드포인트
@app.post("/analyze", summary = "소스 코드 보안 취약점 분석", response_model = AnalyzeResponse)
def analyze(request: CodeRequest = Body(...)):
    # 1. CodeBERT 결과
    bert_result = analyze_code(request.code)

    # 2. GPT 보고서 생성
    gpt_report = generate_report(request.code, label=bert_result["label"])

    # 3. GPT 보고서에서 라벨 추출
    gpt_label = extract_label_from_report(gpt_report)

    # 4. GPT 라벨이 존재하고, CodeBERT 라벨과 다르다면 GPT 결과를 반영
    final = resolve_final_label_score(
        bert_label=bert_result["label"],
        gpt_label=gpt_label
    )

    return {
        "prediction": final["msg"],
        "label": final["label"],
        "security_score": final["score"],
        "report": gpt_report,
        "model_reference": bert_result
    }

# API 분석 보고서 docx 다운로드
@app.post("/report/docx", 
            summary="DOCX 형식의 취약점 보고서 다운로드", 
            description="""
            입력된 소스 코드를 기반으로 CodeBERT와 GPT API를 활용해 보안 취약점을 분석하여
            취약점 종류, 점수, 설명, 공격 시나리오, 보완 방안이 포함된 분석 보고서를
            Word 파일 형식으로 다운로드할 수 있습니다.
            """
)
def download_docx(request: CodeRequest):
    # 1. 모델 예측 및 리포트 생성
    bert_result = analyze_code(request.code)
    gpt_report = generate_report(request.code, label=bert_result["label"])
    gpt_label = extract_label_from_report(gpt_report)

    final = resolve_final_label_score(
        bert_label=bert_result["label"],
        gpt_label=gpt_label
    )
    
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

    # 2. 문서 객체 생성
    doc = Document()
    doc.add_heading("AI 기반 보안 분석 리포트", 0)
    doc.add_paragraph(f"분석 일시: {timestamp}")
    doc.add_paragraph(f"예측된 취약점 라벨: {final['label']}")
    doc.add_paragraph(f"보안 점수: {final['score']}")
    doc.add_paragraph(f"분석 요약 메시지: {final['msg']}")
    doc.add_paragraph("상세 보안 리포트:")
    doc.add_paragraph(gpt_report)

    # 3. 파일 저장
    filename = f"report_{uuid.uuid4().hex[:8]}.docx"
    doc.save(filename)

    return FileResponse(filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="Security_Report.docx")